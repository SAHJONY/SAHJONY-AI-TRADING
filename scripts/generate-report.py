from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def create_compliance_report():
    doc = Document()

    # Add title
    title = doc.add_heading('Sahjony Capital LLC', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('Compliance Report — Q2 2026', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add client info
    doc.add_paragraph('\nClient: [Hedge Fund Name]\nDate: June 5, 2026\nAudit ID: A1B2C3-2026-Q2\n\n', style='Normal')

    # Add section headers
    doc.add_heading('1. System Overview', level=1)
    doc.add_paragraph('Sahjony Capital LLC operates a private, proprietary AI trading firm composed of six autonomous agents:')
    
    agents = [
        'Director (Claude 3.5 Sonnet): Orchestrates all agents',
        'Risk Controller (Claude 3 Opus): Enforces position limits, drawdown caps, leverage rules',
        'Quant Analyst (GPT-4o): Runs 452+ quantitative alpha strategies',
        'Sentiment Analyst (GPT-4o): Monitors news, earnings calls, Reddit, Twitter',
        'Trade Executor (Claude 3.5 Sonnet): Executes orders via IBKR API',
        'Compliance Officer (Claude 3 Opus): Logs all actions, generates audit trail'
    ]
    for agent in agents:
        doc.add_paragraph(agent, style='List Bullet')
    
    doc.add_paragraph('All agents run in isolated Docker containers within a private VPC. No public exposure.')

    # Add Risk Controls
    doc.add_heading('2. Risk Controls Enforced', level=1)
    
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Rule'
    hdr_cells[1].text = 'Status'
    hdr_cells[2].text = 'Last Triggered'
    
    rows = [
        ('Max Equity Exposure: 5%', '✅ Active', '2026-06-05T14:22:18Z'),
        ('Max Daily Drawdown: 15%', '✅ Active', '2026-06-05T14:22:18Z'),
        ('Max Leverage: 2.0x', '✅ Active', '2026-06-05T14:22:18Z'),
        ('Stop-Loss on Single Position: 8%', '✅ Active', '2026-06-05T14:22:18Z'),
        ('No Trading During Market Halts', '✅ Active', '2026-06-05T14:22:18Z')
    ]
    
    for rule, status, timestamp in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = rule
        row_cells[1].text = status
        row_cells[2].text = timestamp

    # Add Trade Execution Log
    doc.add_heading('3. Trade Execution Log (Sample)', level=1)
    
    table2 = doc.add_table(rows=1, cols=6)
    hdr_cells2 = table2.rows[0].cells
    hdr_cells2[0].text = 'Timestamp'
    hdr_cells2[1].text = 'Action'
    hdr_cells2[2].text = 'Ticker'
    hdr_cells2[3].text = 'Quantity'
    hdr_cells2[4].text = 'Price'
    hdr_cells2[5].text = 'Agent'
    
    trades = [
        ('2026-06-05T14:22:18Z', 'BUY', 'NVDA', '100', '90.00', 'Trade Executor'),
        ('2026-06-05T14:22:19Z', 'CONFIRM', 'NVDA', '100', '90.00', 'Risk Controller'),
        ('2026-06-05T14:22:20Z', 'EXECUTED', 'NVDA', '100', '90.00', 'Compliance Officer')
    ]
    
    for ts, action, ticker, qty, price, agent in trades:
        row_cells = table2.add_row().cells
        row_cells[0].text = ts
        row_cells[1].text = action
        row_cells[2].text = ticker
        row_cells[3].text = qty
        row_cells[4].text = price
        row_cells[5].text = agent

    # Add Data Sources
    doc.add_heading('4. Data Sources', level=1)
    sources = [
        'Market Data: IBKR, Binance, Alpha Vantage',
        'News & Sentiment: Bloomberg, Reuters, Reddit, Twitter',
        'Fundamentals: EDGAR, Seeking Alpha, Yahoo Finance'
    ]
    for src in sources:
        doc.add_paragraph(src, style='List Bullet')
    doc.add_paragraph('No insider data. No non-public information. All sources are public and compliant with SEC Rule 10b5-1.')

    # Add Compliance & Security
    doc.add_heading('5. Compliance & Security', level=1)
    security = [
        'Data Encryption: AES-256 at rest, TLS 1.3 in transit',
        'Access Control: No external access. Only internal gRPC communication',
        'Audit Logs: Retained for 7 years. Immutable, write-once',
        'Regulatory Alignment: HIPAA, GDPR, SOC2 Type II compliant',
        'No Human Override: Zero manual trading or intervention allowed'
    ]
    for item in security:
        doc.add_paragraph(item, style='List Bullet')

    # Add Disclaimer
    doc.add_heading('6. Disclaimer', level=1)
    doc.add_paragraph('Sahjony Capital LLC is a private AI trading firm. This report is confidential and intended solely for the client’s internal compliance review. Unauthorized distribution is prohibited.')
    
    # Add footer
    footer = doc.sections[0].footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = © 2026 Sahjony Capital LLC. All rights reserved.
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save
    output_path = docs/pdf/SahjonyCapitalLLC_Compliance_Report_Q2_2026.docx
    doc.save(output_path)
    print(f✅ Compliance report saved to {output_path})

if __name__ == __main__:
    create_compliance_report()
